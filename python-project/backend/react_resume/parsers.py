from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from backend.react_resume.models import ParsedResume


class UnsupportedFileTypeError(ValueError):
    pass


def parse_resume_file(file_name: str, content: bytes) -> ParsedResume:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(file_name, content)
    if suffix == ".docx":
        return _parse_docx(file_name, content)
    if suffix == ".doc":
        raise UnsupportedFileTypeError("暂不支持 .doc，请先转换为 .docx 后上传。")
    raise UnsupportedFileTypeError("仅支持 PDF 和 DOCX 简历文件。")


def _parse_docx(file_name: str, content: bytes) -> ParsedResume:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text.strip())
    text = "\n".join(item for item in [*paragraphs, *table_cells] if item)
    if not text.strip():
        raise ValueError("没有提取到有效文本，请检查简历内容。")
    return ParsedResume(file_name=file_name, file_type="docx", raw_text=text, warnings=[])


def _parse_pdf(file_name: str, content: bytes) -> ParsedResume:
    reader = PdfReader(BytesIO(content))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join(page for page in pages if page)
    warnings: list[str] = []
    if len(text) < 80:
        warnings.append("PDF 可提取文本较少，可能是扫描件或排版复杂。")
    if not text.strip():
        raise ValueError("PDF 没有提取到有效文本，可能是扫描件。")
    return ParsedResume(file_name=file_name, file_type="pdf", raw_text=text, warnings=warnings)
