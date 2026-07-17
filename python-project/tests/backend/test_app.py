from fastapi.testclient import TestClient

from backend.react_resume.models import AgentTrace, OptimizeResumeResponse
from backend.react_resume.app import app


def test_health_check_returns_ok():
    client = TestClient(app)

    response = client.get("/api/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_optimize_resume_requires_supported_file():
    client = TestClient(app)

    response = client.post(
        "/api/optimize-resume",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
        },
        files={"resume_file": ("resume.txt", b"text", "text/plain")},
    )

    assert response.status_code == 400
    assert "仅支持 PDF 和 DOCX" in response.json()["detail"]


def test_optimize_resume_returns_workflow_response(monkeypatch):
    from docx import Document
    from io import BytesIO
    from unittest.mock import Mock

    fake_workflow = Mock()
    fake_workflow.run.return_value = OptimizeResumeResponse(
        diagnosis="技能描述偏泛",
        react_trace=[
            AgentTrace(
                round=1,
                agent="DecisionAgent",
                thought="分析差距",
                action="analyze_resume_gap",
                observation="缺少量化指标",
            )
        ],
        optimized_resume="## 定制简历",
        interview_questions=["Redis 缓存一致性怎么保证？"],
        warnings=[],
    )
    monkeypatch.setattr("backend.react_resume.app.build_workflow", lambda: fake_workflow)

    document = Document()
    document.add_paragraph("熟悉 Java 后端开发，使用 Redis 和 MySQL。")
    stream = BytesIO()
    document.save(stream)

    client = TestClient(app)
    response = client.post(
        "/api/optimize-resume",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
        },
        files={
            "resume_file": (
                "resume.docx",
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["diagnosis"] == "技能描述偏泛"
    assert response.json()["optimized_resume"] == "## 定制简历"


def test_optimize_resume_accepts_pasted_resume_text(monkeypatch):
    from unittest.mock import Mock

    fake_workflow = Mock()
    fake_workflow.run.return_value = OptimizeResumeResponse(
        diagnosis="技能描述偏泛",
        react_trace=[
            AgentTrace(
                round=1,
                agent="DecisionAgent",
                thought="分析差距",
                action="analyze_resume_gap",
                observation="缺少量化指标",
            )
        ],
        optimized_resume="## 定制简历",
        interview_questions=["Redis 缓存一致性怎么保证？"],
        warnings=[],
    )
    monkeypatch.setattr("backend.react_resume.app.build_workflow", lambda: fake_workflow)

    client = TestClient(app)
    response = client.post(
        "/api/optimize-resume",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
            "resume_text": "张三\n熟悉 Java 后端开发，使用 Redis 和 MySQL 做过缓存优化。",
        },
    )

    assert response.status_code == 200
    assert response.json()["optimized_resume"] == "## 定制简历"
    parsed_resume = fake_workflow.run.call_args.kwargs["parsed_resume"]
    assert parsed_resume.file_type == "text"
    assert "熟悉 Java 后端开发" in parsed_resume.raw_text


def test_optimize_resume_requires_file_or_pasted_text():
    client = TestClient(app)

    response = client.post(
        "/api/optimize-resume",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "请上传简历文件或粘贴简历内容。"


def test_optimize_resume_runs_workflow_in_threadpool(monkeypatch):
    from docx import Document
    from io import BytesIO
    from unittest.mock import Mock

    calls = []

    async def fake_run_in_threadpool(func, *args, **kwargs):
        calls.append(func.__name__)
        return func(*args, **kwargs)

    fake_workflow = Mock()
    fake_workflow.run.return_value = OptimizeResumeResponse(
        diagnosis="技能描述偏泛",
        react_trace=[
            AgentTrace(
                round=1,
                agent="DecisionAgent",
                thought="分析差距",
                action="analyze_resume_gap",
                observation="缺少量化指标",
            )
        ],
        optimized_resume="## 定制简历",
        interview_questions=["Redis 缓存一致性怎么保证？"],
        warnings=[],
    )
    monkeypatch.setattr("backend.react_resume.app.run_in_threadpool", fake_run_in_threadpool)
    monkeypatch.setattr("backend.react_resume.app.build_workflow", lambda: fake_workflow)

    document = Document()
    document.add_paragraph("熟悉 Java 后端开发，使用 Redis 和 MySQL。")
    stream = BytesIO()
    document.save(stream)

    client = TestClient(app)
    response = client.post(
        "/api/optimize-resume",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
        },
        files={
            "resume_file": (
                "resume.docx",
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert calls == ["process_resume_request"]


def test_optimize_resume_stream_returns_sse_events(monkeypatch):
    class FakeWorkflow:
        def stream(self, **kwargs):
            yield {"type": "status", "message": "正在诊断匹配差距"}
            yield {"type": "answer_delta", "content": "## 定制简历\n面向后端开发岗位"}
            yield {
                "type": "result",
                "data": OptimizeResumeResponse(
                    diagnosis="技能描述偏泛",
                    react_trace=[],
                    optimized_resume="## 定制简历",
                    interview_questions=["Redis 缓存一致性怎么保证？"],
                    warnings=[],
                ).model_dump(),
            }
            yield {"type": "done"}

    monkeypatch.setattr("backend.react_resume.app.build_workflow", lambda: FakeWorkflow())

    client = TestClient(app)
    response = client.post(
        "/api/optimize-resume/stream",
        data={
            "target_company": "字节跳动",
            "target_role": "后端开发",
            "jd_text": "要求 Java 和 Redis",
            "resume_text": "张三\n熟悉 Java 后端开发，使用 Redis 和 MySQL 做过缓存优化。",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    assert "event: status" in response.text
    assert '"type": "result"' in response.text
    assert "定制简历" in response.text
